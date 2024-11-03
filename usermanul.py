
from datetime import date, datetime
import os
import uuid
import streamlit as st
import fitz  # PyMuPDF
import nltk
from PIL import Image
import io
import base64
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from cryptography.fernet import Fernet
import re
from streamlit_chat import message
from docx import Document
from docx.shared import Inches
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import concurrent.futures
import streamlit as st
import asyncio

# Directory containing the PDFs
current_file_directory = os.path.dirname(os.path.abspath(__file__))
relative_path = "MARUTHI"
base_directory = os.path.join(current_file_directory, relative_path)

# Download NLTK data
nltk.download("punkt")
stemmer = PorterStemmer()


async def correct_spelling_with_google_genai(text):
    try:
        model = genai.GenerativeModel(
            "gemini-1.5-flash",
            safety_settings={
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
            },
            generation_config=genai.GenerationConfig(temperature=0),
        )
        response = model.generate_content(
            contents=f"Correct the spelling errors in the following text. Do not add any additional text or sentences:\n\n{text}",
        )
        corrected_text = response.text
    except Exception as e:
        st.error(f"An error occurred while correcting spelling: {e}")
        corrected_text = text

    return corrected_text


def convert_latex_to_text(latex_text):
    superscript_map = {
        "0": "⁰",
        "1": "¹",
        "2": "²",
        "3": "³",
        "4": "⁴",
        "5": "⁵",
        "6": "⁶",
        "7": "⁷",
        "8": "⁸",
        "9": "⁹",
    }
    subscript_map = {
        "0": "₀",
        "1": "₁",
        "2": "₂",
        "3": "₃",
        "4": "₄",
        "5": "₅",
        "6": "₆",
        "7": "₇",
        "8": "₈",
        "9": "₉",
    }

    def replace_superscripts(match):
        base = match.group(1)
        superscript = match.group(2)
        return base + "".join(superscript_map.get(ch, ch) for ch in superscript)

    def replace_subscripts(match):
        base = match.group(1)
        subscript = match.group(2)
        return base + "".join(subscript_map.get(ch, ch) for ch in subscript)

    latex_text = re.sub(r"\^([0-9]+)\^", replace_superscripts, latex_text)
    latex_text = re.sub(r"~([0-9]+)~", replace_subscripts, latex_text)

    return latex_text


def load_api_key():
    with open("secret.key", "rb") as key_file:
        key = key_file.read()
    fernet = Fernet(key)
    with open("encrypted_api_key.txt", "rb") as encrypted_file:
        encrypted_api_key = encrypted_file.read()
    decrypted_api_key = fernet.decrypt(encrypted_api_key).decode()
    return decrypted_api_key


api_key = load_api_key()
genai.configure(api_key=api_key)


async def extract_text_and_images(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text("text")
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            images.append(resize_image(image))
    return text, images


def load_pdf_data(pdf_path):
    with concurrent.futures.ThreadPoolExecutor() as executor:
        future = executor.submit(extract_text_and_images, pdf_path)
        text, images = future.result()  # Expect only two values
    return text, images

 
def preprocess_text(text):
    """Tokenize and stem the text."""
    tokens = word_tokenize(text.lower())
    return " ".join(stemmer.stem(token) for token in tokens if token.isalpha())


def search_text_in_pdf(query, text):
    preprocessed_query = preprocess_text(query)

    # Tokenize the text and preprocess
    sentences = nltk.sent_tokenize(text)
    preprocessed_sentences = [preprocess_text(sentence) for sentence in sentences]

    # Use TF-IDF Vectorizer to transform text
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([preprocessed_query] + preprocessed_sentences)

    # Compute cosine similarity
    cosine_sim = cosine_similarity(vectors[0:1], vectors[1:])
    relevant_indices = cosine_sim[0].argsort()[-5:][
        ::-1
    ]  # Get top 5 relevant sentences

    relevant_sentences = [sentences[i] for i in relevant_indices]
    return " ".join(relevant_sentences)


def resize_image(image, max_width=900, max_height=600):
    original_width, original_height = image.size
    aspect_ratio = original_width / original_height
    new_width, new_height = original_width, original_height
    if original_width > max_width or original_height > max_height:
        if aspect_ratio > 1:
            new_width = max_width
            new_height = int(max_width / aspect_ratio)
        else:
            new_height = max_height
            new_width = int(max_height * aspect_ratio)
        image = image.resize((new_width, new_height), Image.LANCZOS)
    return image


def display_images(images):
    html = '<div style="display: flex; flex-wrap: wrap; gap: 10px;">'
    for img in images:
        img = resize_image(img)
        buffered = io.BytesIO()
        img.save(buffered, format="PNG")
        img_data = base64.b64encode(buffered.getvalue()).decode("utf-8")
        html += f'<img src="data:image/png;base64,{img_data}" style="border-radius: 15px; max-width: 100%; height: auto; margin: 5px;" />'
    html += "</div>"
    return html


async def refine_text_with_google_genai(query, result_text):
    try:
        model = genai.GenerativeModel(
            "gemini-1.5-flash",
            safety_settings={
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
            },
            generation_config=genai.GenerationConfig(temperature=2),
        )
        response = model.generate_content(
            contents=f"Refine the following text to focus on the user query '{query}' and if derivation contain show each step and remove any unrelated information:\n\n{result_text}",
        )
        refined_text = response.text
    except Exception as e:
        st.error(f"An error occurred while refining the text: {e}")
        refined_text = result_text

    superscript_map = {
        "^0": "⁰",
        "^1": "¹",
        "^2": "²",
        "^3": "³",
        "^4": "⁴",
        "^5": "⁵",
        "^6": "⁶",
        "^7": "⁷",
        "^8": "⁸",
        "^9": "⁹",
        "^x": "ˣ",
        "^+": "⁺",
        "^-": "⁻",
        "^=": "⁼",
        "^(": "⁽",
        "^)": "⁾",
    }
    subscript_map = {
        "~0~": "₀",
        "~1~": "₁",
        "~2~": "₂",
        "~3~": "₃",
        "~4~": "₄",
        "~5~": "₅",
        "~6~": "₆",
        "~7~": "₇",
        "~8~": "₈",
        "~9~": "₉",
        "~+~": "₊",
        "~−~": "₋",
        "~=~": "₌",
        "~(~": "₍",
        "~)~": "₎",
    }

    for key, value in superscript_map.items():
        refined_text = refined_text.replace(key, value)
    for key, value in subscript_map.items():
        refined_text = refined_text.replace(key, value)
    refined_text = refined_text.replace("µ", "µ").replace("λ", "λ")

    return refined_text

 
def clean_message_content(content):
    # Remove markdown formatting
    content = re.sub(r"\*\*(.*?)\*\*", r"\1", content)
    content = re.sub(r"(.*?)", r"\1", content)  # Remove text
    content = re.sub(r"## ", "", content)
    # Remove "Relevant Images/Diagrams:" text and any HTML
    content = re.sub(
        r"Relevant Images/Diagrams:\s*<div.*?</div>", "", content, flags=re.DOTALL
    )

    return content


def save_chat_as_word():
    docs = Document()
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            docs.add_heading(text="• " + msg["content"], level=1)
        else:
            clean_content = clean_message_content(msg["content"])
            docs.add_paragraph(text=clean_content, style="BodyText")

        if "Relevant Images/Diagrams:" in msg["content"]:
            images = re.findall(
                r"data:image/png;base64,([A-Za-z0-9+/=]+)", msg["content"]
            )
            for img_data in images:
                imgbyte = base64.b64decode(img_data)
                imgstream = io.BytesIO(imgbyte)
                docs.add_picture(imgstream, width=Inches(2.0))

    doc_bytes = io.BytesIO()
    docs.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


async def main():
    st.title("KTUAssistant")
    st.text("Your personal assistant bot for engineering study materials.")
    if "step" not in st.session_state:
        st.session_state.step = 1
    if "selected_stream" not in st.session_state:
        st.session_state.selected_stream = ""
    if "selected_year" not in st.session_state:
        st.session_state.selected_year = 1
    if "subject" not in st.session_state:
        st.session_state.subject = ""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "loading" not in st.session_state:
        st.session_state.loading = False

async def find_relevant_images(query, pdf_files, all_image_info):
    relevant_images = []
    seen_images = set()
    query = (
            query.lower().strip()
        )  # Normalize query: lower case and remove extra spaces

    for pdf_file in pdf_files:
            pdf_path = os.path.join(subject_directory, pdf_file)
            doc = fitz.open(pdf_path)

            for page_num, img_index in all_image_info:
                if page_num < len(doc):
                    page = doc.load_page(page_num)
                    page_text = page.get_text(
                        "text"
                    ).lower()  # Convert page text to lowercase

                    # Check if the query is present in the page text
                    if query in page_text:
                        image_list = page.get_images(full=True)
                        if img_index < len(image_list):
                            base_image = doc.extract_image(image_list[img_index][0])
                            image_bytes = base_image["image"]
                            image = Image.open(io.BytesIO(image_bytes))

                            image_hash = hash(image_bytes)
                            if image_hash not in seen_images:
                                seen_images.add(image_hash)

 
if (
                                    not (image.width == 500 and image.height == 500)
                                    and not (image.width == 273 and image.height == 95)
                                    and not (image.width == 471 and image.height == 494)
                                    and not (
                                        image.width == 3264 and image.height == 3264
                                    )
                                    and not (
                                        image.width == 1595 and image.height == 267
                                    )
                                    and not (image.width == 932 and image.height == 67)
                                    and not (image.width == 468 and image.height == 99)
                                    and not (image.width == 109 and image.height == 23)
                                    and not (image.width == 112 and image.height == 126)
                                    and not (image.width == 96 and image.height == 84)
                                    and not (image.width == 252 and image.height == 284)
                                    and not (image.width == 146 and image.height == 164)
                                    and not (
                                        image.width == 1293 and image.height == 290
                                    )
                                    and not (
                                        image.width == 1291 and image.height == 219
                                    )
                                    and not (image.width == 947 and image.height == 269)
                                    and not (
                                        image.width == 1242 and image.height == 161
                                    )
                                    and not (
                                        image.width == 1344 and image.height == 329
                                    )
                                    and not (image.width == 808 and image.height == 401)
                                    and not (
                                        image.width == 1344 and image.height == 318
                                    )
                                    and not (
                                        image.width == 1341 and image.height == 212
                                    )
                                    and not (image.width == 370 and image.height == 193)
                                    and not (image.width == 503 and image.height == 83)
                                ):
                                    relevant_images.append(image)

        return relevant_images

    def back():
        if st.session_state.step == 3:
            st.session_state.step = 2
        elif st.session_state.step == 2:
            st.session_state.step = 1

    def next():
        if st.session_state.step == 1:
            st.session_state.selected_stream = selected_stream
            st.session_state.step = 2
        elif st.session_state.step == 2:
            st.session_state.selected_year = selected_year
            st.session_state.step = 3
        elif st.session_state.step == 3:
            st.session_state.subject = selected_subject
            st.session_state.step = 4

    if st.session_state.step == 1:
        st.subheader("Select Your Engineering Branch")
        streams = [
            "Computer Science Engineering",
            "Electrical And Electronics Engineering",
            "Mechanical Engineering",
            "Civil Engineering",
            "Electronics And Communication Engineering",
        ]
        selected_stream = st.selectbox("Choose a branch", streams)
        if st.button("Next", key="nextsem", on_click=next):
            pass

 
elif st.session_state.step == 2:
        st.subheader(f"Select Your Semester")
        selected_year = st.selectbox(
            "Choose a semester", ["s1", "s2", "s3", "s4", "s5", "s6", "s7", "s8"]
        )
        col1, col2, col3 = st.columns([0.13, 0.17, 1])
        with col1:
            if st.button("Back", on_click=back):
                pass
        with col2:
            if st.button("Next", key="nextsub", on_click=next):
                pass

    elif st.session_state.step == 3:
        st.subheader(f"Select Your Course")
        s1 = [
            "Engineering Physics A - (PHT100)",
            "Engineering Chemistry - (CYT100)",
            "Engineering Mechanics - (EST100)",
            "Engineering Graphics - (EST110)",
            "Basics Of Civil & Mechanical Engineering - (EST120)",
            "Basics Of Electrical & Electronics Engineering - (EST130)",
            "Engineering Physics Lab - (PHL120)",
            "Engineering Chemistry Lab - (CYL120)",
            "Civil & Mechanical Workshop - (ESL120)",
            "Electrical & Electronics Workshop - (ESL130)",
            "Life Skills - (HUN101)",
            "Linear Algebra And Calculus - (MAT101)",
        ]
        s2 = [
            "Vector Calculus, Differential Equations And Transforms - (MAT102)",
            "Engineering Physics A - (PHT100)",
            "Engineering Chemistry - (CYT100)",
            "Engineering Mechanics - (EST100)",
            "Engineering Graphics - (EST110)",
            "Basics Of Civil & Mechanical Engineering - (EST120)",
            "Basics Of Electrical & Electronics Engineering - (EST130)",
            "Professional Communication - (HUN102)",
            "Programming In C - (EST102)",
            "Engineering Physics Lab - (PHL120)",
            "Engineering Chemistry Lab - (CYL120)",
            "Civil & Mechanical Workshop - (ESL120)",
            "Electrical & Electronics Workshop - (ESL130)",
        ]
        css3 = [
            "Discrete Mathematical Structures - (MAT203)",
            "Data Structures - (CST201)",
            "Logic System Design - (CST203)",
            "Object Oriented Programming Using Java - (CST205)",
            "Design And Engineering - (EST200)",
            "Professional Ethics - (HUT200)",
            "Sustainable Engineering - (MCN201)",
        ]
        css4 = [
            "Graph Theory - (MAT206)",
            "Computer Organisation And Architecture - (CST202)",
            "Database Management Systems - (CST204)",
            "Operating Systems - (CST206)",
            "Design And Engineering - (EST200)",
            "Professional Ethics - (HUT200)",
            "Constitution Of India - (MCN202)",
            "Digital Lab - (CSL202)",
            "Operating Systems Lab - (CSL204)",
        ]
        css5 = [
            "Formal Languages And Automata Theory - (CST301)",
            "Computer Networks - (CST303)",
            "System Software - (CST305)",
            "Microprocessors And Microcontrollers - (CST307)",
            "Management Of Software Systems - (CST309)",
            "Disaster Management - (MCN301)",
            "System Software And Microprocessors Lab - (CSL331)",
            "Database Management Systems Lab - (CSL333)",
        ]
        css6 = [
            "Compiler Design - (CST302)",
            "Computer Graphics And Image Processing - (CST304)",
            "Algorithm Analysis And Design - (CST306)",
            "Elective - I",
            "Industrial Economics And Foreign Trade - (HUT300)",
            "Comprehensive Course Work - (CST308)",
            "Networking Lab - (CSL332)",
            "Miniproject - (CSD334)",
        ]
        css7 = [
            "Artificial Intelligence - (CST401)",
            "Elective - II",
            "Elective - Open",
            "Industrial Safety Engineering - (MCN401)",
            "Compiler Lab - (CSL411)",
            "Seminar - (CSQ413)",
            "Project Phase I - (CSD415)",
        ]
        css8 = [
            "Distributed Computing - (CST402)",
            "Elective - III",

 
"Elective - IV",
            "Elective - V",
            "Comprehensive Course Viva - (CST404)",
            "Project Phase II - (CSD416)",
        ]
        ecs3 = [
            "Partial Differential Equation And Complex Analysis - (MAT201)",
            "Solid State Devices - (ECT201)",
            "Logic Circuit Design - (ECT203)",
            "Network Theory - (ECT205)",
            "Design And Engineering - (EST200)",
            "Sustainable Engineering - (MCN201)",
            "Logic Design Lab - (ECL203)",
            "Scientific Computing Lab - (ECL201)",
        ]
        ecs4 = [
            "Probability, Random Processes And Numerical Methods - (MAT204)",
            "Analog Circuits - (ECT202)",
            "Signals And Systems - (ECT204)",
            "Computer Architecture And Microcontrollers - (ECT206)",
            "Design And Engineering - (EST200)",
            "Professional Ethics - (HUT200)",
            "Constitution Of India - (MCN202)",
            "Analog Circuits And Simulation Lab - (ECL202)",
            "Microcontroller Lab - (ECL204)",
        ]
        ecs5 = [
            "Linear Integrated Circuits - (ECT301)",
            "Digital Signal Processing - (ECT303)",
            "Analog And Digital Communication - (ECT305)",
            "Control Systems - (ECT307)",
            "Industrial Economics And Foreign Trade - (HUT300)",
            "Management For Engineers - (HUT310)",
            "Disaster Management - (MCN301)",
            "Analog Integrated Circuits And Simulation Lab - (ECL331)",
            "Digital Signal Processing Lab - (ECL333)",
        ]
        ecs6 = [
            "Electromagnetics - (ECT302)",
            "VlSI Circuit Design - (ECT304)",
            "Information Theory And Coding - (ECT306)",
            "Elective - I",
            "Industrial Economics And Foreign Trade - (HUT300)",
            "Management For Engineers - (HUT310)",
            "Comprehensive Course Work - (ECT308)",
            "Communication Lab - (ECL332)",
            "Miniproject - (ECD334)",
        ]
        ecs7 = [
            "Wireless Communication - (ECT401)",
            "Elective - II",
            "Elective - Open",
            "Industrial Safety Engineering - (MCN401)",
            "Electromagnetics Lab - (ECL411)",
            "Seminar - (ECQ413)",
            "Project Phase I - (ECD415)",
        ]
        ecs8 = [
            "Instrumentation - (ECT402)",
            "Elective - III",
            "Elective - IV",
            "Elective - V",
            "Comprehensive Viva Voce - (ECT404)",
            "Project Phase II - (ECD416)",
        ]

 
if (
            st.session_state.selected_stream == "Computer Science Engineering"
            or st.session_state.selected_stream
            == "Electrical And Electronics Engineering"
            or st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            or st.session_state.selected_stream == "Civil Engineering"
            or st.session_state.selected_stream == "Mechanical Engineering"
        ) and st.session_state.selected_year == "s1":
            selected_subject = st.selectbox("choose a course", s1)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            or st.session_state.selected_stream
            == "Electrical And Electronics Engineering"
            or st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            or st.session_state.selected_stream == "Civil Engineering"
            or st.session_state.selected_stream == "Mechanical Engineering"
        ) and st.session_state.selected_year == "s2":
            selected_subject = st.selectbox("choose a course", s2)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s3"
        ):
            selected_subject = st.selectbox("choose a course", css3)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s4"
        ):
            selected_subject = st.selectbox("choose a course", css4)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s5"
        ):
            selected_subject = st.selectbox("choose a course", css5)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s6"
        ):
            selected_subject = st.selectbox("choose a course", css6)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s7"
        ):
            selected_subject = st.selectbox("choose a course", css7)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):

 
pass
        elif (
            st.session_state.selected_stream == "Computer Science Engineering"
            and st.session_state.selected_year == "s8"
        ):
            selected_subject = st.selectbox("choose a course", css8)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass

        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s3"
        ):
            selected_subject = st.selectbox("choose a course", ecs3)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s4"
        ):
            selected_subject = st.selectbox("choose a course", ecs4)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s5"
        ):
            selected_subject = st.selectbox("choose a course", ecs5)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s6"
        ):
            selected_subject = st.selectbox("choose a course", ecs6)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s7"
        ):
            selected_subject = st.selectbox("choose a course", ecs7)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass
        elif (
            st.session_state.selected_stream
            == "Electronics And Communication Engineering"
            and st.session_state.selected_year == "s8"
        ):
            selected_subject = st.selectbox("choose a course", ecs8)
            col1, col2, col3 = st.columns([0.13, 0.17, 1])
            with col1:
                if st.button("Back", on_click=back):
                    pass
            with col2:
                if st.button("Submit", on_click=next):
                    pass

 
elif st.session_state.step == 4:
        st.subheader(
            f"\n{st.session_state.selected_stream} - {st.session_state.selected_year}\n{st.session_state.subject}"
        )
        st.divider()
        stream_directory = os.path.join(
            base_directory, st.session_state.selected_stream
        )
        year_directory = os.path.join(stream_directory, st.session_state.selected_year)
        subject_directory = os.path.join(year_directory, st.session_state.subject)
        pdf_files = [f for f in os.listdir(subject_directory) if f.endswith(".pdf")]
        aggregated_text = ""
        all_images = []
        all_image_info = []

        for pdf_file in pdf_files:
            pdf_path = os.path.join(subject_directory, pdf_file)

            text, images = await extract_text_and_images(pdf_path)

            aggregated_text += text
            all_images.extend(images)

        text_col, button_col, backcol = st.columns(
            [4, 0.58, 0.9], vertical_alignment="bottom"
        )

        # Define styling for both elements
        # Define styling for text input and button
        style = """
        <style>
          .stTextInput {
            position: fixed;
            bottom: 60px;
            z-index: 9999;
            
          }

          .stButton {
            position: fixed;
            bottom: 60px;
            z-index: 9999;
          }
          
          .stbackbtn{
            position: fixed;
            bottom: 60px;
            z-index: 9999;
          }
          
          .fixed-text {
            position: fixed;
            bottom: 20px;
            z-index: 9999;
            padding-left:220px;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.2);
            }
        </style>
        """

        # Inject the styling code for both elements
        st.markdown(style, unsafe_allow_html=True)

        def newchat():
            st.session_state.step = 1
            st.session_state.messages = []

        # Add your text input and button
        with text_col:
            textcrt = st.text_input("Enter Your Topic Heading:", key="stTextInput")
            user_query = await correct_spelling_with_google_genai(text=textcrt)
        with button_col:
            btn = st.button("Send", key="stButton")
        with backcol:
            if st.button("New Chat", key="stbackbtn", on_click=newchat):
                pass
        st.markdown(
            '<div class="fixed-text">KTUASSISTANT can make mistakes.</div>',
            unsafe_allow_html=True,
        )
        if btn:
            if user_query:
                st.session_state.messages.append(
                    {"role": "user", "content": user_query}
                )
                st.session_state.loading = True

                # Show loading indicator
                with st.spinner("Bot is processing your request..."):
                    result_text = search_text_in_pdf(user_query, aggregated_text)
                    if result_text:
                        refined_text = await refine_text_with_google_genai(
                            user_query, result_text
                        )
                        formatted_text = convert_latex_to_text(refined_text)
                        response_message = f"\n\n{formatted_text}\n"
                    else:
                        response_message = "No relevant notes found."

                    relevant_images = await find_relevant_images(
                        user_query, pdf_files, all_image_info
                    )
                    if relevant_images:
                        images_html = display_images(relevant_images)
                        response_message += (
                            "\n\nRelevant Images/Diagrams:\n" + images_html
                        )
                    else:
                        response_message += "\n\nNo relevant images/diagrams found."

 
# Update chat history with bot response and hide loading indicator
                st.session_state.messages.append(
                    {"role": "bot", "content": response_message}
                )
                st.session_state.loading = False

        # Display chat history
        for msg in st.session_state.messages:
            unique_key = str(uuid.uuid4())
            if msg["role"] == "user":
                message(msg["content"], is_user=True, key=unique_key)
            else:
                message(msg["content"], is_user=False, allow_html=True, key=unique_key)

        if st.session_state.loading:
            st.spinner("Bot is processing your request...")
        times = datetime.now()
        formatted = times.strftime("%Y%m%d%H%M%S")
        if btn:
            st.divider()
        if btn:
            doc_bytes = save_chat_as_word()
            st.download_button(
                label="Download",
                data=doc_bytes,
                file_name=f"ktuassistant{formatted}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


if name == "main":
    asyncio.run(main())